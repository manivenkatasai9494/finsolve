import os
import pandas as pd
from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from pinecone import Pinecone, ServerlessSpec
from langchain_pinecone import PineconeVectorStore
from docx import Document as DocxDocument
from pypdf import PdfReader

load_dotenv()

splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=100)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")

# ---------- Universal File Reader ----------
def read_file(file_path):
    if file_path.endswith((".txt", ".md")):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    elif file_path.endswith(".csv"):
        return pd.read_csv(file_path).to_string()

    elif file_path.endswith(".xlsx"):
        return pd.read_excel(file_path).to_string()

    elif file_path.endswith(".docx"):
        doc = DocxDocument(file_path)
        return "\n".join([p.text for p in doc.paragraphs])

    elif file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        return "".join([p.extract_text() or "" for p in reader.pages])

    return ""

# ---------- RBAC ----------
def get_roles_by_domain(domain):
    domain = domain.lower()

    if domain == "engineering":
        return ["Engineering", "C-Level"]
    elif domain == "finance":
        return ["Finance", "C-Level"]
    elif domain == "hr":
        return ["HR", "C-Level"]
    elif domain == "marketing":
        return ["Marketing", "C-Level"]
    elif domain == "general":
        return ["Employee", "HR", "Finance", "Marketing", "Engineering", "C-Level"]

    return ["C-Level"]

# ---------- Load Data ----------
def load_documents():
    docs = []

    for domain in os.listdir(DATA_DIR):
        path = os.path.join(DATA_DIR, domain)

        if not os.path.isdir(path):
            continue

        for file in os.listdir(path):
            file_path = os.path.join(path, file)
            text = read_file(file_path)

            chunks = splitter.split_text(text)

            for chunk in chunks:
                docs.append(
                    Document(
                        page_content=chunk,
                        metadata={
                            "domain": domain,
                            "allowed_roles": get_roles_by_domain(domain)
                        }
                    )
                )

    return docs

# ---------- Build Index ----------
def build_index():
    documents = load_documents()

    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )

    pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
    index_name = os.getenv("PINECONE_INDEX", "company-rag")

    if index_name in [i.name for i in pc.list_indexes()]:
        pc.delete_index(index_name)

    pc.create_index(
        name=index_name,
        dimension=384,
        metric="cosine",
        spec=ServerlessSpec(cloud="aws", region="us-east-1")
    )

    PineconeVectorStore.from_documents(
        documents,
        embedding=embeddings,
        index_name=index_name
    )

    print("Index created successfully")

if __name__ == "__main__":
    build_index()